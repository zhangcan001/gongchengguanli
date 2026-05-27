import re
import sqlite3
from datetime import date
from pathlib import Path
from typing import Any
from uuid import uuid4

from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter

from .archive_service import ArchiveService
from .config import Settings
from .errors import ErrorCode
from .issues import IssueService
from .progress_analytics import ProgressAnalyticsService
from .repositories import RepositoryError


FILE_ASSET_COLUMNS = (
    "id",
    "project_id",
    "business_type",
    "business_id",
    "file_name",
    "original_file_name",
    "file_path",
    "file_type",
    "mime_type",
    "file_size",
    "uploaded_by",
    "uploaded_at",
)

DIARY_FIELDS = (
    ("weather", "天气"),
    ("temperature", "温度"),
    ("construction_summary", "今日施工情况"),
    ("workers_summary", "施工人员情况"),
    ("machinery_summary", "施工机械情况"),
    ("quality_summary", "质量检查情况"),
    ("safety_summary", "安全检查情况"),
    ("patrol_summary", "巡视检查情况"),
    ("issue_summary", "存在问题"),
    ("handling_opinion", "处理意见"),
    ("tomorrow_plan", "明日重点"),
)

ISSUE_TYPE_LABELS = {
    "quality": "质量问题",
    "safety": "安全隐患",
    "progress": "进度滞后",
    "document": "资料缺失",
    "drawing": "图纸问题",
    "other": "其他问题",
}

ISSUE_LEVEL_LABELS = {
    "normal": "普通",
    "important": "重要",
    "urgent": "紧急",
    "major": "重大",
}

ISSUE_STATUS_LABELS = {
    "pending_rectification": "待整改",
    "notified": "已通知",
    "replied": "已回复",
    "pending_review": "待复查",
    "closed": "已关闭",
    "archived": "已归档",
    "overdue": "已逾期",
    "rejected": "已驳回",
    "reopened": "重新打开",
}

DELAY_LEVEL_LABELS = {
    "normal_or_ahead": "正常或超前",
    "slight_delay": "轻微滞后",
    "obvious_delay": "明显滞后",
    "serious_delay": "严重滞后",
}

WORD_MIME_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
EXCEL_MIME_TYPE = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"


def _ensure_project_exists(connection: sqlite3.Connection, project_id: int) -> None:
    row = connection.execute("SELECT 1 FROM project WHERE id = ?", (project_id,)).fetchone()
    if not row:
        raise RepositoryError(ErrorCode.PROJECT_NOT_FOUND, "Project not found.")


def _row_to_dict(row: sqlite3.Row) -> dict[str, Any]:
    return {column: row[column] for column in row.keys()}


def _asset_to_dict(row: sqlite3.Row) -> dict[str, Any]:
    item = {column: row[column] for column in FILE_ASSET_COLUMNS}
    item["download_url"] = f"/api/files/{item['id']}/download"
    return item


def _blank(value: Any, fallback: str = "未填写") -> str:
    if value is None:
        return fallback
    text = str(value).strip()
    return text if text else fallback


def _date_text(value: Any) -> str:
    return _blank(value, date.today().isoformat())


def _location(*parts: Any) -> str:
    text = "".join(str(part).strip() for part in parts if part)
    return text or "项目"


def _sanitize_filename(value: str, *, max_length: int = 120) -> str:
    cleaned = re.sub(r'[<>:"/\\|?*\x00-\x1f]', "_", value).strip(" .")
    cleaned = re.sub(r"\s+", " ", cleaned)
    if not cleaned:
        cleaned = "export"
    return cleaned[:max_length].strip(" .") or "export"


def _build_original_name(*, day: Any, location: str, topic: str, doc_type: str, suffix: str) -> str:
    stem = "_".join(
        _sanitize_filename(str(part), max_length=48)
        for part in (_date_text(day), location, topic, doc_type)
        if _blank(part, "") != ""
    )
    return f"{stem}.{suffix.lstrip('.')}"


def _mime_type_for_suffix(suffix: str) -> str:
    normalized = suffix.lower().lstrip(".")
    if normalized == "docx":
        return WORD_MIME_TYPE
    if normalized == "xlsx":
        return EXCEL_MIME_TYPE
    return "application/octet-stream"


def _label(mapping: dict[str, str], value: Any) -> str:
    return mapping.get(str(value or ""), _blank(value))


class ExportService:
    def __init__(self, connection: sqlite3.Connection, settings: Settings) -> None:
        self.connection = connection
        self.settings = settings

    def export_diary_word(self, diary_id: int) -> dict[str, Any]:
        diary = self._get_diary(diary_id)
        document = self._new_document("监理日志")
        self._add_meta(document, diary["project_name"], f"日志日期：{diary['diary_date']}")
        self._add_key_value_table(
            document,
            [
                ("项目名称", diary["project_name"]),
                ("日志日期", diary["diary_date"]),
                ("天气", diary["weather"]),
                ("温度", diary["temperature"]),
                ("确认状态", "已确认" if diary["confirmed"] else "未确认"),
            ],
        )
        for key, label in DIARY_FIELDS[2:]:
            self._add_section(document, label, diary[key])

        original_name = _build_original_name(
            day=diary["diary_date"],
            location="项目",
            topic="监理日志",
            doc_type="监理日志",
            suffix="docx",
        )
        return self._save_document(
            document,
            project_id=diary["project_id"],
            business_type="diary_export",
            business_id=diary_id,
            original_file_name=original_name,
        )

    def export_patrol_word(self, patrol_id: int) -> dict[str, Any]:
        patrol = self._get_patrol(patrol_id)
        document = self._new_document("巡视记录")
        location = _location(patrol["building"], patrol["floor"], patrol["area"])
        self._add_meta(document, patrol["project_name"], f"巡视日期：{patrol['patrol_date']}")
        self._add_key_value_table(
            document,
            [
                ("项目名称", patrol["project_name"]),
                ("巡视日期", patrol["patrol_date"]),
                ("巡视人员", patrol["patrol_person"]),
                ("巡视部位", location),
                ("专业", patrol["discipline"]),
                ("是否生成问题", "是" if patrol["generate_issue"] else "否"),
                ("关联问题", f"#{patrol['issue_id']}" if patrol["issue_id"] else "无"),
            ],
        )
        self._add_section(document, "巡视内容", patrol["content"])
        self._add_section(document, "发现问题", patrol["found_problem"])
        self._add_section(document, "处理意见", patrol["handling_opinion"])

        original_name = _build_original_name(
            day=patrol["patrol_date"],
            location=location,
            topic="巡视记录",
            doc_type="巡视记录",
            suffix="docx",
        )
        return self._save_document(
            document,
            project_id=patrol["project_id"],
            business_type="patrol_export",
            business_id=patrol_id,
            original_file_name=original_name,
        )

    def export_issue_notice_word(self, issue_id: int) -> dict[str, Any]:
        issue = IssueService(self.connection).get_issue(issue_id)
        document = self._new_document("整改通知单")
        location = _location(issue.get("building"), issue.get("floor"), issue.get("area"))
        self._add_meta(document, issue.get("project_name"), f"问题编号：#{issue_id}")
        self._add_key_value_table(
            document,
            [
                ("项目名称", issue.get("project_name")),
                ("问题类型", _label(ISSUE_TYPE_LABELS, issue.get("issue_type"))),
                ("问题等级", _label(ISSUE_LEVEL_LABELS, issue.get("level"))),
                ("责任单位", issue.get("responsible_unit")),
                ("发现日期", issue.get("discovered_date")),
                ("整改期限", issue.get("deadline")),
                ("问题部位", location),
                ("发现人", issue.get("discovered_by")),
            ],
        )
        self._add_section(document, "问题描述", issue.get("description"))
        self._add_section(document, "整改要求", issue.get("rectification_requirement"))
        self._add_section(document, "监理意见", "请责任单位按要求限期完成整改，并报监理复查。")

        original_name = _build_original_name(
            day=issue.get("discovered_date"),
            location=location,
            topic=issue.get("title") or "整改通知",
            doc_type="整改通知单",
            suffix="docx",
        )
        return self._save_document(
            document,
            project_id=issue["project_id"],
            business_type="issue_notice_export",
            business_id=issue_id,
            original_file_name=original_name,
        )

    def export_issue_review_word(self, issue_id: int) -> dict[str, Any]:
        issue = IssueService(self.connection).get_issue(issue_id)
        document = self._new_document("整改复查记录")
        location = _location(issue.get("building"), issue.get("floor"), issue.get("area"))
        review_actions = [action for action in issue["actions"] if action["action_type"] in {"review", "close"}]
        latest_review = review_actions[-1] if review_actions else None
        self._add_meta(document, issue.get("project_name"), f"问题编号：#{issue_id}")
        self._add_key_value_table(
            document,
            [
                ("项目名称", issue.get("project_name")),
                ("问题标题", issue.get("title")),
                ("问题部位", location),
                ("责任单位", issue.get("responsible_unit")),
                ("当前状态", _label(ISSUE_STATUS_LABELS, issue.get("effective_status"))),
                ("关闭时间", issue.get("closed_at")),
                ("复查日期", latest_review.get("action_date") if latest_review else None),
                ("复查人", latest_review.get("operator") if latest_review else None),
            ],
        )
        self._add_section(document, "整改要求", issue.get("rectification_requirement"))
        self._add_section(document, "复查意见", latest_review.get("content") if latest_review else "暂无复查记录。")

        if review_actions:
            document.add_heading("复查流转记录", level=2)
            table = document.add_table(rows=1, cols=4)
            table.style = "Table Grid"
            headers = ("动作", "日期", "操作人", "内容")
            for index, header in enumerate(headers):
                table.rows[0].cells[index].text = header
            for action in review_actions:
                cells = table.add_row().cells
                cells[0].text = "关闭" if action["action_type"] == "close" else "复查"
                cells[1].text = _blank(action.get("action_date"))
                cells[2].text = _blank(action.get("operator"))
                cells[3].text = _blank(action.get("content"), "")

        original_name = _build_original_name(
            day=(latest_review or {}).get("action_date") or issue.get("closed_at") or date.today().isoformat(),
            location=location,
            topic=issue.get("title") or "整改复查",
            doc_type="整改复查记录",
            suffix="docx",
        )
        return self._save_document(
            document,
            project_id=issue["project_id"],
            business_type="issue_review_export",
            business_id=issue_id,
            original_file_name=original_name,
        )

    def export_issues_excel(self, project_id: int) -> dict[str, Any]:
        _ensure_project_exists(self.connection, project_id)
        project = self._get_project(project_id)
        issues = IssueService(self.connection).list_issues(project_id=project_id)
        workbook = Workbook()
        worksheet = workbook.active
        worksheet.title = "问题台账"
        rows = [
            [
                "编号",
                "类型",
                "等级",
                "状态",
                "标题",
                "部位",
                "专业",
                "责任单位",
                "发现人",
                "发现日期",
                "整改期限",
                "关闭时间",
                "整改要求",
            ]
        ]
        for issue in issues:
            rows.append(
                [
                    issue["id"],
                    _label(ISSUE_TYPE_LABELS, issue["issue_type"]),
                    _label(ISSUE_LEVEL_LABELS, issue["level"]),
                    _label(ISSUE_STATUS_LABELS, issue["effective_status"]),
                    issue["title"],
                    _location(issue.get("building"), issue.get("floor"), issue.get("area")),
                    issue.get("discipline") or "",
                    issue.get("responsible_unit") or "",
                    issue.get("discovered_by") or "",
                    issue.get("discovered_date") or "",
                    issue.get("deadline") or "",
                    issue.get("closed_at") or "",
                    issue.get("rectification_requirement") or "",
                ]
            )
        self._write_sheet_table(worksheet, rows)
        original_name = _build_original_name(
            day=date.today().isoformat(),
            location=project["name"],
            topic="问题台账",
            doc_type="问题台账",
            suffix="xlsx",
        )
        return self._save_workbook(
            workbook,
            project_id=project_id,
            business_type="issue_ledger_export",
            business_id=None,
            original_file_name=original_name,
        )

    def export_progress_analysis_excel(self, project_id: int) -> dict[str, Any]:
        _ensure_project_exists(self.connection, project_id)
        project = self._get_project(project_id)
        analytics = ProgressAnalyticsService(self.connection)
        overview = analytics.get_overview(project_id)
        delay = analytics.get_delay_analysis(project_id)
        quality = analytics.get_data_quality(project_id)

        workbook = Workbook()
        overview_sheet = workbook.active
        overview_sheet.title = "进度概览"
        self._write_sheet_table(
            overview_sheet,
            [
                ["指标", "值"],
                ["项目名称", project["name"]],
                ["最新数据日期", overview["latest_data_date"] or ""],
                ["总体实际完成率", overview["overall_actual_percent"] if overview["overall_actual_percent"] is not None else ""],
                ["总体计划完成率", overview["overall_planned_percent"] if overview["overall_planned_percent"] is not None else ""],
                ["偏差", overview["deviation"] if overview["deviation"] is not None else ""],
                ["滞后等级", _label(DELAY_LEVEL_LABELS, overview["delay_level"]) if overview["delay_level"] else "无法判断"],
                ["无可计算进度", "是" if overview["no_calculable_progress"] else "否"],
                ["最近导入批次", f"#{overview['latest_batch']['id']} {overview['latest_batch']['file_name']}" if overview["latest_batch"] else ""],
            ],
        )
        self._write_sheet_table(
            workbook.create_sheet("楼栋统计"),
            [["楼栋", "实际完成率", "计划完成率", "偏差", "滞后等级", "记录数"]]
            + [
                [
                    item["label"],
                    item["actual_percent"] if item["actual_percent"] is not None else "",
                    item["planned_percent"] if item["planned_percent"] is not None else "",
                    item["deviation"] if item["deviation"] is not None else "",
                    _label(DELAY_LEVEL_LABELS, item["delay_level"]) if item["delay_level"] else "",
                    item["record_count"],
                ]
                for item in overview["building_summary"]
            ],
        )
        self._write_sheet_table(
            workbook.create_sheet("专业统计"),
            [["专业", "实际完成率", "计划完成率", "偏差", "滞后等级", "记录数"]]
            + [
                [
                    item["label"],
                    item["actual_percent"] if item["actual_percent"] is not None else "",
                    item["planned_percent"] if item["planned_percent"] is not None else "",
                    item["deviation"] if item["deviation"] is not None else "",
                    _label(DELAY_LEVEL_LABELS, item["delay_level"]) if item["delay_level"] else "",
                    item["record_count"],
                ]
                for item in overview["discipline_summary"]
            ],
        )
        self._write_sheet_table(
            workbook.create_sheet("滞后任务"),
            [["任务", "楼栋", "楼层", "专业", "计划完成率", "实际完成率", "偏差", "滞后等级", "备注"]]
            + [
                [
                    task.get("task_name") or "",
                    task.get("building") or "",
                    task.get("floor") or "",
                    task.get("discipline") or "",
                    task.get("planned_percent") if task.get("planned_percent") is not None else "",
                    task.get("actual_percent") if task.get("actual_percent") is not None else "",
                    task.get("deviation"),
                    _label(DELAY_LEVEL_LABELS, task.get("delay_level")),
                    task.get("remark") or "",
                ]
                for task in delay["delayed_tasks"]
            ],
        )
        self._write_sheet_table(
            workbook.create_sheet("数据质量"),
            [["级别", "字段", "任务", "楼栋", "楼层", "专业", "数据日期", "提示"]]
            + [
                [
                    item["severity"],
                    item["field"],
                    item.get("task_name") or "",
                    item.get("building") or "",
                    item.get("floor") or "",
                    item.get("discipline") or "",
                    item.get("data_date") or "",
                    item["message"],
                ]
                for item in [*quality["error_items"], *quality["warning_items"]]
            ],
        )

        original_name = _build_original_name(
            day=date.today().isoformat(),
            location=project["name"],
            topic="进度分析",
            doc_type="进度分析",
            suffix="xlsx",
        )
        return self._save_workbook(
            workbook,
            project_id=project_id,
            business_type="progress_analysis_export",
            business_id=overview["latest_batch"]["id"] if overview["latest_batch"] else None,
            original_file_name=original_name,
        )

    def get_file_asset(self, file_id: int) -> dict[str, Any]:
        row = self.connection.execute(
            f"SELECT {', '.join(FILE_ASSET_COLUMNS)} FROM file_asset WHERE id = ?",
            (file_id,),
        ).fetchone()
        if not row:
            raise RepositoryError(ErrorCode.FILE_ASSET_NOT_FOUND, "File asset not found.")
        return _asset_to_dict(row)

    def resolve_download_path(self, file_id: int) -> tuple[Path, dict[str, Any]]:
        asset = self.get_file_asset(file_id)
        base_dir = self.settings.data_dir.resolve()
        file_path = (base_dir / asset["file_path"]).resolve()
        try:
            file_path.relative_to(base_dir)
        except ValueError as exc:
            raise RepositoryError(ErrorCode.FILE_NOT_FOUND, "File path is outside data directory.") from exc
        if not file_path.is_file():
            raise RepositoryError(ErrorCode.FILE_NOT_FOUND, "File not found on disk.")
        return file_path, asset

    def _get_project(self, project_id: int) -> dict[str, Any]:
        row = self.connection.execute("SELECT * FROM project WHERE id = ?", (project_id,)).fetchone()
        if not row:
            raise RepositoryError(ErrorCode.PROJECT_NOT_FOUND, "Project not found.")
        return _row_to_dict(row)

    def _get_diary(self, diary_id: int) -> dict[str, Any]:
        row = self.connection.execute(
            """
            SELECT diary.*, project.name AS project_name
            FROM diary
            LEFT JOIN project ON project.id = diary.project_id
            WHERE diary.id = ?
            """,
            (diary_id,),
        ).fetchone()
        if not row:
            raise RepositoryError(ErrorCode.DIARY_NOT_FOUND, "Diary not found.")
        item = _row_to_dict(row)
        item["confirmed"] = bool(item["confirmed"])
        item["ai_generated"] = bool(item["ai_generated"])
        return item

    def _get_patrol(self, patrol_id: int) -> dict[str, Any]:
        row = self.connection.execute(
            """
            SELECT patrol_record.*, project.name AS project_name
            FROM patrol_record
            LEFT JOIN project ON project.id = patrol_record.project_id
            WHERE patrol_record.id = ?
            """,
            (patrol_id,),
        ).fetchone()
        if not row:
            raise RepositoryError(ErrorCode.PATROL_RECORD_NOT_FOUND, "Patrol record not found.")
        item = _row_to_dict(row)
        item["generate_issue"] = bool(item["generate_issue"])
        item["write_to_diary"] = bool(item["write_to_diary"])
        return item

    def _new_document(self, title: str) -> DocumentObject:
        document = Document()
        styles = document.styles
        styles["Normal"].font.name = "Microsoft YaHei"
        styles["Normal"].font.size = Pt(10.5)
        heading = document.add_heading(title, level=0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return document

    def _add_meta(self, document: DocumentObject, project_name: Any, subtitle: str) -> None:
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(f"{_blank(project_name)} | {subtitle}")
        run.font.size = Pt(10)
        run.italic = True
        document.add_paragraph("")

    def _add_key_value_table(self, document: DocumentObject, rows: list[tuple[str, Any]]) -> None:
        table = document.add_table(rows=0, cols=2)
        table.style = "Table Grid"
        for label, value in rows:
            cells = table.add_row().cells
            cells[0].text = label
            cells[1].text = _blank(value)
        document.add_paragraph("")

    def _add_section(self, document: DocumentObject, title: str, content: Any) -> None:
        document.add_heading(title, level=2)
        paragraph = document.add_paragraph(_blank(content, "暂无记录。"))
        paragraph.paragraph_format.first_line_indent = Pt(21)

    def _save_document(
        self,
        document: DocumentObject,
        *,
        project_id: int,
        business_type: str,
        business_id: int | None,
        original_file_name: str,
    ) -> dict[str, Any]:
        target_path = self._export_path(original_file_name)
        document.save(target_path)
        return self._insert_file_asset(
            project_id=project_id,
            business_type=business_type,
            business_id=business_id,
            target_path=target_path,
            original_file_name=original_file_name,
        )

    def _save_workbook(
        self,
        workbook: Workbook,
        *,
        project_id: int,
        business_type: str,
        business_id: int | None,
        original_file_name: str,
    ) -> dict[str, Any]:
        target_path = self._export_path(original_file_name)
        workbook.save(target_path)
        return self._insert_file_asset(
            project_id=project_id,
            business_type=business_type,
            business_id=business_id,
            target_path=target_path,
            original_file_name=original_file_name,
        )

    def _export_path(self, original_file_name: str) -> Path:
        export_dir = self.settings.data_dir / "files" / "exports"
        export_dir.mkdir(parents=True, exist_ok=True)
        suffix = Path(original_file_name).suffix
        stem = _sanitize_filename(Path(original_file_name).stem, max_length=90)
        return export_dir / f"{uuid4().hex}_{stem}{suffix}"

    def _insert_file_asset(
        self,
        *,
        project_id: int,
        business_type: str,
        business_id: int | None,
        target_path: Path,
        original_file_name: str,
    ) -> dict[str, Any]:
        file_size = target_path.stat().st_size
        file_name = target_path.name
        suffix = target_path.suffix.lower().lstrip(".")
        cursor = self.connection.execute(
            """
            INSERT INTO file_asset (
                project_id,
                business_type,
                business_id,
                file_name,
                original_file_name,
                file_path,
                file_type,
                mime_type,
                file_size,
                uploaded_by
            )
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (
                project_id,
                business_type,
                business_id,
                file_name,
                original_file_name,
                str(target_path.relative_to(self.settings.data_dir)),
                suffix,
                _mime_type_for_suffix(suffix),
                file_size,
                "system",
            ),
        )
        self.connection.commit()
        asset = self.get_file_asset(int(cursor.lastrowid))
        archive = ArchiveService(self.connection, self.settings).archive_file_asset(
            file_id=asset["id"],
            business_type=business_type,
            business_id=business_id,
        )
        asset["archive_id"] = archive["id"]
        asset["archive_path"] = archive["archive_path"]
        return asset

    def _write_sheet_table(self, worksheet: Any, rows: list[list[Any]]) -> None:
        if not rows:
            rows = [["暂无数据"]]
        for row in rows:
            worksheet.append(row)
        header_fill = PatternFill("solid", fgColor="1E5BFF")
        header_font = Font(color="FFFFFF", bold=True)
        for cell in worksheet[1]:
            cell.fill = header_fill
            cell.font = header_font
            cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        worksheet.freeze_panes = "A2"
        for column_cells in worksheet.columns:
            column_letter = get_column_letter(column_cells[0].column)
            max_length = max(len(str(cell.value or "")) for cell in column_cells)
            worksheet.column_dimensions[column_letter].width = min(max(max_length + 4, 12), 36)
            for cell in column_cells:
                cell.alignment = Alignment(vertical="top", wrap_text=True)

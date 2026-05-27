import sqlite3
from datetime import date

from docx import Document
from openpyxl import load_workbook

from tests.test_issues import create_issue
from tests.test_progress_import import upload_progress_excel
from tests.test_smart_inbox import create_project


def _asset_path(client, payload):
    return client.app.state.settings.data_dir / payload["file_path"]


def _assert_export_asset(client, payload, *, suffix: str, business_type: str):
    assert payload["file_id"] if "file_id" in payload else payload["id"]
    assert payload["business_type"] == business_type
    assert payload["file_type"] == suffix
    assert payload["download_url"] == f"/api/files/{payload['id']}/download"
    path = _asset_path(client, payload)
    assert path.is_file()
    assert path.stat().st_size == payload["file_size"]

    with sqlite3.connect(client.app.state.settings.database_path) as connection:
        row = connection.execute("SELECT business_type, file_path FROM file_asset WHERE id = ?", (payload["id"],)).fetchone()
    assert row[0] == business_type
    assert row[1] == payload["file_path"]
    return path


def _create_confirmed_diary(client, project_id: int):
    generated = client.post(
        "/api/diary/generate",
        json={
            "project_id": project_id,
            "diary_date": "2026-05-26",
            "weather": "晴",
            "temperature": "25-32℃",
            "manual_note": "今日现场施工正常。",
        },
    ).json()
    response = client.post(
        "/api/diary/confirm",
        json={
            "project_id": project_id,
            "diary_date": "2026-05-26",
            "weather": "晴",
            "temperature": "25-32℃",
            "ai_generation_id": generated["ai_generation_id"],
            "draft": generated["draft"],
        },
    )
    assert response.status_code == 200
    return response.json()


def _create_patrol(client, project_id: int):
    response = client.post(
        "/api/quick-record/confirm",
        json={
            "project_id": project_id,
            "confirmed_fields": {
                "building": "3#楼",
                "floor": "12层",
                "area": "砌体",
                "discipline": "土建",
                "issue_type": "quality",
                "description": "砌体灰缝不饱满",
                "patrol_content": "3#楼12层砌体灰缝不饱满，要求施工单位整改。",
                "rectification_requirement": "请施工单位整改后报监理复查。",
                "patrol_person": "王监理",
            },
            "confirmed_actions": ["create_patrol"],
        },
    )
    assert response.status_code == 200
    return response.json()["patrol_record_id"]


def _publish_progress(client, project_id: int):
    upload = upload_progress_excel(client, project_id)
    assert upload.status_code == 200
    batch = client.post(
        "/api/progress/import/analyze",
        json={"project_id": project_id, "inbox_id": upload.json()["inbox_id"]},
    ).json()
    response = client.post(f"/api/progress/import/{batch['batch_id']}/publish", json={"replace_existing": False})
    assert response.status_code == 200
    return batch["batch_id"]


def test_diary_word_export_creates_file_asset_and_download(client):
    project = create_project(client)
    diary = _create_confirmed_diary(client, project["id"])

    response = client.post(f"/api/diary/{diary['id']}/export")

    assert response.status_code == 200
    payload = response.json()
    path = _assert_export_asset(client, payload, suffix="docx", business_type="diary_export")
    document = Document(path)
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert "监理日志" in text
    assert "今日施工情况" in text

    download = client.get(payload["download_url"])
    assert download.status_code == 200
    assert len(download.content) == payload["file_size"]


def test_patrol_word_export(client):
    project = create_project(client)
    patrol_id = _create_patrol(client, project["id"])

    response = client.post(f"/api/patrol/{patrol_id}/export")

    assert response.status_code == 200
    payload = response.json()
    path = _assert_export_asset(client, payload, suffix="docx", business_type="patrol_export")
    document = Document(path)
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert "巡视记录" in text
    assert "巡视内容" in text


def test_issue_notice_word_export(client):
    project = create_project(client)
    issue = create_issue(client, project["id"])

    response = client.post(f"/api/issues/{issue['id']}/export-notice")

    assert response.status_code == 200
    payload = response.json()
    path = _assert_export_asset(client, payload, suffix="docx", business_type="issue_notice_export")
    document = Document(path)
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert "整改通知单" in text
    assert "整改要求" in text


def test_issue_review_word_export(client):
    project = create_project(client)
    issue = create_issue(client, project["id"])
    client.post(f"/api/issues/{issue['id']}/notify", json={"content": "通知整改", "operator": "王监理"})
    client.post(f"/api/issues/{issue['id']}/reply", json={"content": "已整改", "operator": "施工单位"})
    client.post(f"/api/issues/{issue['id']}/close", json={"content": "复查合格", "operator": "王监理"})

    response = client.post(f"/api/issues/{issue['id']}/export-review")

    assert response.status_code == 200
    payload = response.json()
    path = _assert_export_asset(client, payload, suffix="docx", business_type="issue_review_export")
    document = Document(path)
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert "整改复查记录" in text
    assert "复查意见" in text


def test_issues_excel_export(client):
    project = create_project(client)
    issue = create_issue(client, project["id"])

    response = client.post(f"/api/issues/export-excel?project_id={project['id']}")

    assert response.status_code == 200
    payload = response.json()
    path = _assert_export_asset(client, payload, suffix="xlsx", business_type="issue_ledger_export")
    workbook = load_workbook(path)
    worksheet = workbook["问题台账"]
    assert worksheet["A1"].value == "编号"
    assert worksheet["A2"].value == issue["id"]
    assert "问题台账" in payload["original_file_name"]


def test_progress_analysis_excel_export(client):
    project = create_project(client)
    batch_id = _publish_progress(client, project["id"])

    response = client.post("/api/progress/export-analysis", json={"project_id": project["id"]})

    assert response.status_code == 200
    payload = response.json()
    path = _assert_export_asset(client, payload, suffix="xlsx", business_type="progress_analysis_export")
    workbook = load_workbook(path)
    assert {"进度概览", "楼栋统计", "专业统计", "滞后任务", "数据质量"}.issubset(set(workbook.sheetnames))
    assert workbook["进度概览"]["B2"].value == project["name"]
    assert payload["business_id"] == batch_id


def test_file_download_reports_missing_asset(client):
    response = client.get("/api/files/999/download")

    assert response.status_code == 404
    assert response.json()["detail"]["code"] == "FILE_ASSET_NOT_FOUND"
